"""
INFORME EDIFICIO PARTE 2 OK 1
"""

#!/usr/bin/env python
# coding: utf-8

# In[2]:


import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import matplotlib.pyplot as plt

def read_excel_data(file_path, sheet_name):
    """Lee los datos de la hoja de Excel especificada y los devuelve como un DataFrame."""
    with pd.ExcelFile(file_path) as excel_data:
        return pd.read_excel(excel_data, sheet_name=sheet_name)

def desdinamizar_data(data, id_vars):
    """Desdinamiza los datos (convierte de formato ancho a formato largo)."""
    value_vars = data.columns.difference(id_vars)  # Columnas de los meses
    data_largo = pd.melt(data, id_vars=id_vars, value_vars=value_vars, 
                         var_name='Mes', value_name='Monto')
    return data_largo

def ordenar_por_mes(data):
    """Ordena los datos por la columna 'Mes' utilizando el orden cronológico de los meses."""
    meses_orden = ['enero', 'febrero', 'marzo', 'abril', 'mayo', 'junio', 
                   'julio', 'agosto', 'septiembre', 'octubre', 'noviembre', 'diciembre']
    data['Mes'] = data['Mes'].str.lower().str.strip()
    data['Mes'] = pd.Categorical(data['Mes'], categories=meses_orden, ordered=True)
    data.sort_values(by='Mes', inplace=True)
    return data

def create_word_report(departments_data, expensas_data, agua_data, output_path, selected_departments):
    """Crea un informe Word con los depósitos agrupados por departamento, incluyendo las expensas y consumo de agua."""
    doc = Document()

    for department in selected_departments:
        if department in departments_data.groups:
            data = departments_data.get_group(department)
        else:
            data = pd.DataFrame()  # Crear un DataFrame vacío si no hay datos del departamento, valor de depósitos es 0

        add_department_to_document(doc, department, data, expensas_data, agua_data)

    doc.save(output_path)
    print(f"Informe generado y guardado en: {output_path}")

def add_department_to_document(doc, department, data, expensas_data, agua_data):
    """Añade la información de un departamento, sus expensas, consumo de agua y el balance de pagos al documento Word."""
    doc.add_heading(f'Departamento {department}', level=1)

    # Añadir el subtítulo de "Pagos realizados en favor del edificio"
    doc.add_heading(f"Pagos realizados en favor del edificio", level=2)

    # Añadir la información de depósitos
    if not data.empty:
        for index, row in data.iterrows():
            monto_rounded = round(row['Monto'], 2)
            doc.add_paragraph(f"Fecha: {row['Fecha']}, Hora: {row['Hora']}, Monto: {monto_rounded}, Nota: {row['Nota']}")
    else:
        doc.add_paragraph("No se encontraron pagos para este departamento.")

    # Obtener las expensas y consumo de agua del departamento
    expensas_dept = expensas_data[expensas_data['Departamento'] == department].copy()
    agua_dept = agua_data[agua_data['Departamento'] == department].copy()

    # Ordenar por mes utilizando el orden cronológico
    expensas_dept = ordenar_por_mes(expensas_dept)
    agua_dept = ordenar_por_mes(agua_dept)

    # Agregar subtítulo de expensas
    doc.add_heading(f"Expensas desde junio hasta el mes actual", level=2)

    # Crear tabla para expensas
    if not expensas_dept.empty:
        table_expensas = doc.add_table(rows=1, cols=2)
        hdr_cells = table_expensas.rows[0].cells
        hdr_cells[0].text = 'Mes'
        hdr_cells[1].text = 'Monto Expensas'

        for index, row in expensas_dept.iterrows():
            monto_rounded = round(row['Monto'], 2)
            row_cells = table_expensas.add_row().cells
            row_cells[0].text = str(row['Mes'])
            row_cells[1].text = str(monto_rounded)
    else:
        doc.add_paragraph("No se encontraron expensas para este departamento.")

    # Agregar subtítulo de consumo de agua
    doc.add_heading(f"Costo de consumo agua desde junio hasta el mes actual", level=2)

    # Crear tabla para consumo de agua
    if not agua_dept.empty:
        table_agua = doc.add_table(rows=1, cols=2)
        hdr_cells = table_agua.rows[0].cells
        hdr_cells[0].text = 'Mes'
        hdr_cells[1].text = 'Monto Agua'

        for index, row in agua_dept.iterrows():
            monto_rounded = round(row['Monto'], 2)
            row_cells = table_agua.add_row().cells
            row_cells[0].text = str(row['Mes'])
            row_cells[1].text = str(monto_rounded)
    else:
        doc.add_paragraph("No se encontró información de agua para este departamento.")

    # Añadir el subtítulo de "Balance de pagos"
    doc.add_heading(f"Balance de pagos", level=2)

    # Generar gráfico de balance de pagos
    total_monto = round(data['Monto'].sum(), 2) if not data.empty else 0
    total_expensas_agua = round(expensas_dept['Monto'].sum() + agua_dept['Monto'].sum(), 2)
    generate_balance_chart(department, total_monto, total_expensas_agua, doc)
    generate_difference_phrase(department, total_monto, total_expensas_agua, doc)

    doc.add_page_break()

def generate_balance_chart(department, total_monto, total_expensas_agua, doc):
    """Genera un gráfico de barras horizontales con el balance de pagos y lo inserta en el documento Word."""
    fig, ax = plt.subplots()

    bars = ax.barh(['Pagos realizados', 'Expensas y Agua'], [total_monto, total_expensas_agua], 
                   color=['blue', 'red'], label=['Pagos realizados', 'Expensas y Agua'])

    for bar in bars:
        ax.text(bar.get_width(), bar.get_y() + bar.get_height()/2, f'{bar.get_width():,.2f}', 
                va='center', ha='left')

    ax.set_xlabel('Monto')
    ax.set_title(f'Balance de pagos del departamento {department}')
    ax.legend()

    image_path = f'balance_{department}.png'
    plt.savefig(image_path, bbox_inches='tight')
    plt.close()

    doc.add_picture(image_path, width=Inches(5))

def generate_difference_phrase(department, total_monto, total_expensas_agua, doc):
    """Genera una frase basada en la diferencia entre pagos realizados y expensas + agua."""
    diferencia = round(total_monto - total_expensas_agua, 2)

    paragraph = doc.add_paragraph()
    if diferencia < 0:
        run = paragraph.add_run(f"Este departamento/tienda debe al edificio {abs(diferencia):,.2f}Bs.")
    else:
        run = paragraph.add_run(f"Le felicitamos, usted está al día y tiene a su favor el monto de {diferencia:,.2f}Bs.")

    run.bold = True
    run.font.size = Pt(14)

def main():
    # Ruta del archivo Excel y del informe Word
    file_path = r"C:\Users\HP\Desktop\EDIFICIO JUAN BOSCO\EXPENSAS\EXPENSAS CON CONCILIACION BANCARIA\CONCILIACION BANCARIA COPIA SEGURA.xlsx"
    output_path = r'C:\Users\HP\Desktop\EDIFICIO JUAN BOSCO\EXPENSAS\EXPENSAS CON CONCILIACION BANCARIA\Informe_por departamento.docx'

    # Leer los datos de las hojas 'INFORME BANCARIO', 'EXPENSAS' y 'COSTO AGUA DEPAS'
    df_informe_bancario = read_excel_data(file_path, 'INFORME BANCARIO')
    df_expensas = read_excel_data(file_path, 'EXPENSAS')
    df_agua = read_excel_data(file_path, 'COSTO AGUA DEPAS')

    # Desdinamizar los datos de expensas y agua (formato largo)
    df_expensas_largo = desdinamizar_data(df_expensas, id_vars=['Departamento'])
    df_agua_largo = desdinamizar_data(df_agua, id_vars=['Departamento'])

    # Agrupar los datos de los depósitos por departamento
    departments_data = df_informe_bancario.groupby('Departamento')

    # Lista de departamentos seleccionados

    selected_departments = [
        'T5', 'T4', 'T3', 'T2', 'T1',
        '9E', '9D', '9C', '9B', '9A',
        '8E', '8D', '8C', '8B', '8A',
        '7E', '7D', '7C', '7B', '7A',
        '6E', '6D', '6C', '6B', '6A',
        '5E', '5D', '5C', '5B', '5A',
        '4E', '4D', '4C', '4B', '4A',
        '3E', '3D', '3C', '3B', '3A',
        '2E', '2D', '2C', '2B', '2A',
        '1E', '1D', '1C', '1B', '1A']


    # Crear el informe en formato Word solo para los departamentos seleccionados
    create_word_report(departments_data, df_expensas_largo, df_agua_largo, output_path, selected_departments)

# Ejecutar la función principal
if __name__ == "__main__":
    main()


# In[ ]:




